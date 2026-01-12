"""Transcript API endpoints."""

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse, Response
from pydantic import BaseModel
from typing import List, Optional
from datetime import datetime
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from sqlalchemy.orm import selectinload
import os
import io
import json

from src.database.connection import get_db
from src.database.models import Recording, Project, Transcript, TranscriptSegment
from src.api.auth import get_current_user, User

router = APIRouter()


# Pydantic models
class WordTimestamp(BaseModel):
    word: str
    start_time: float
    end_time: float
    confidence: float


class TranscriptSegmentResponse(BaseModel):
    id: str
    index: int
    speaker_id: str
    text: str
    start_time: int  # milliseconds
    end_time: int  # milliseconds
    confidence: float
    words: List[WordTimestamp]
    is_edited: bool

    class Config:
        from_attributes = True


class SpeakerResponse(BaseModel):
    id: str
    name: Optional[str]
    color: str


class TranscriptResponse(BaseModel):
    id: str
    recording_id: str
    language: str
    segments: List[TranscriptSegmentResponse]
    speakers: List[SpeakerResponse]
    model: str
    model_version: str
    processing_time: int
    word_count: int
    character_count: int
    average_confidence: float
    created_at: str
    updated_at: str


class SegmentUpdate(BaseModel):
    text: str


class SpeakerUpdate(BaseModel):
    name: str


class ExportRequest(BaseModel):
    format: str  # txt, json, srt, vtt, docx, pdf
    include_speaker_labels: bool = True
    include_timestamps: bool = True


# Speaker colors for visualization
SPEAKER_COLORS = [
    "#3B82F6",  # blue
    "#22C55E",  # green
    "#A855F7",  # purple
    "#F97316",  # orange
    "#EC4899",  # pink
    "#06B6D4",  # cyan
    "#EAB308",  # yellow
    "#EF4444",  # red
]


def get_speaker_color(index: int) -> str:
    """Get a color for a speaker based on index."""
    return SPEAKER_COLORS[index % len(SPEAKER_COLORS)]


async def verify_recording_access(
    db: AsyncSession, recording_id: str, user_id: str
) -> Recording:
    """Verify user has access to the recording."""
    result = await db.execute(
        select(Recording)
        .options(selectinload(Recording.project))
        .where(Recording.id == recording_id)
    )
    recording = result.scalar_one_or_none()

    if not recording:
        raise HTTPException(status_code=404, detail="Recording not found")

    if recording.project.owner_id != user_id:
        raise HTTPException(status_code=403, detail="Access denied")

    return recording


@router.get("/{transcript_id}/segments/{segment_id}")
async def get_segment(
    transcript_id: str,
    segment_id: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Get a single transcript segment."""
    result = await db.execute(
        select(TranscriptSegment)
        .options(selectinload(TranscriptSegment.transcript).selectinload(Transcript.recording).selectinload(Recording.project))
        .where(TranscriptSegment.id == segment_id, TranscriptSegment.transcript_id == transcript_id)
    )
    segment = result.scalar_one_or_none()

    if not segment:
        raise HTTPException(status_code=404, detail="Segment not found")

    # Verify access
    if segment.transcript.recording.project.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Access denied")

    words = segment.words_json or []

    return TranscriptSegmentResponse(
        id=segment.id,
        index=segment.segment_index,
        speaker_id=segment.speaker_id or "speaker_0",
        text=segment.text,
        start_time=int(segment.start_time * 1000),
        end_time=int(segment.end_time * 1000),
        confidence=segment.confidence,
        words=[
            WordTimestamp(
                word=w.get("word", ""),
                start_time=w.get("start_time", 0),
                end_time=w.get("end_time", 0),
                confidence=w.get("confidence", 0),
            )
            for w in words
        ],
        is_edited=segment.updated_at > segment.created_at,
    )


@router.patch("/{transcript_id}/segments/{segment_id}")
async def update_segment(
    transcript_id: str,
    segment_id: str,
    updates: SegmentUpdate,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Update transcript segment text."""
    result = await db.execute(
        select(TranscriptSegment)
        .options(selectinload(TranscriptSegment.transcript).selectinload(Transcript.recording).selectinload(Recording.project))
        .where(TranscriptSegment.id == segment_id, TranscriptSegment.transcript_id == transcript_id)
    )
    segment = result.scalar_one_or_none()

    if not segment:
        raise HTTPException(status_code=404, detail="Segment not found")

    # Verify access
    if segment.transcript.recording.project.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Access denied")

    # Update segment
    segment.text = updates.text
    segment.updated_at = datetime.utcnow()

    # Update transcript full text and word count
    transcript = segment.transcript
    all_segments = await db.execute(
        select(TranscriptSegment)
        .where(TranscriptSegment.transcript_id == transcript_id)
        .order_by(TranscriptSegment.segment_index)
    )
    segments = all_segments.scalars().all()
    transcript.full_text = " ".join(s.text for s in segments)
    transcript.word_count = len(transcript.full_text.split())
    transcript.updated_at = datetime.utcnow()

    await db.commit()

    words = segment.words_json or []

    return TranscriptSegmentResponse(
        id=segment.id,
        index=segment.segment_index,
        speaker_id=segment.speaker_id or "speaker_0",
        text=segment.text,
        start_time=int(segment.start_time * 1000),
        end_time=int(segment.end_time * 1000),
        confidence=segment.confidence,
        words=[
            WordTimestamp(
                word=w.get("word", ""),
                start_time=w.get("start_time", 0),
                end_time=w.get("end_time", 0),
                confidence=w.get("confidence", 0),
            )
            for w in words
        ],
        is_edited=True,
    )


@router.patch("/{transcript_id}/speakers/{speaker_id}")
async def update_speaker(
    transcript_id: str,
    speaker_id: str,
    updates: SpeakerUpdate,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Update speaker name for all segments with this speaker ID."""
    # Get transcript and verify access
    result = await db.execute(
        select(Transcript)
        .options(selectinload(Transcript.recording).selectinload(Recording.project))
        .where(Transcript.id == transcript_id)
    )
    transcript = result.scalar_one_or_none()

    if not transcript:
        raise HTTPException(status_code=404, detail="Transcript not found")

    if transcript.recording.project.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Access denied")

    # Update all segments with this speaker ID
    segments_result = await db.execute(
        select(TranscriptSegment).where(
            TranscriptSegment.transcript_id == transcript_id,
            TranscriptSegment.speaker_id == speaker_id,
        )
    )
    segments = segments_result.scalars().all()

    for segment in segments:
        segment.speaker_name = updates.name
        segment.updated_at = datetime.utcnow()

    await db.commit()

    return SpeakerResponse(
        id=speaker_id,
        name=updates.name,
        color=get_speaker_color(int(speaker_id.replace("speaker_", "")) if speaker_id.startswith("speaker_") else 0),
    )


def format_srt_time(seconds: float) -> str:
    """Format seconds to SRT time format (HH:MM:SS,mmm)."""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    millis = int((seconds % 1) * 1000)
    return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"


def format_vtt_time(seconds: float) -> str:
    """Format seconds to VTT time format (HH:MM:SS.mmm)."""
    return format_srt_time(seconds).replace(",", ".")


def export_to_txt(segments: List[TranscriptSegment], include_speakers: bool, include_timestamps: bool) -> str:
    """Export transcript to plain text."""
    lines = []
    for seg in segments:
        prefix = ""
        if include_timestamps:
            prefix += f"[{format_vtt_time(seg.start_time)}] "
        if include_speakers and seg.speaker_id:
            name = seg.speaker_name or seg.speaker_id.replace("speaker_", "Speaker ")
            prefix += f"{name}: "
        lines.append(f"{prefix}{seg.text}")
    return "\n\n".join(lines)


def export_to_srt(segments: List[TranscriptSegment], include_speakers: bool) -> str:
    """Export transcript to SRT format."""
    lines = []
    for idx, seg in enumerate(segments, 1):
        lines.append(str(idx))
        lines.append(f"{format_srt_time(seg.start_time)} --> {format_srt_time(seg.end_time)}")
        text = seg.text
        if include_speakers and seg.speaker_id:
            name = seg.speaker_name or seg.speaker_id.replace("speaker_", "Speaker ")
            text = f"[{name}] {text}"
        lines.append(text)
        lines.append("")
    return "\n".join(lines)


def export_to_vtt(segments: List[TranscriptSegment], include_speakers: bool) -> str:
    """Export transcript to WebVTT format."""
    lines = ["WEBVTT", ""]
    for idx, seg in enumerate(segments, 1):
        lines.append(str(idx))
        lines.append(f"{format_vtt_time(seg.start_time)} --> {format_vtt_time(seg.end_time)}")
        text = seg.text
        if include_speakers and seg.speaker_id:
            name = seg.speaker_name or seg.speaker_id.replace("speaker_", "Speaker ")
            text = f"<v {name}>{text}"
        lines.append(text)
        lines.append("")
    return "\n".join(lines)


def export_to_json(transcript: Transcript, segments: List[TranscriptSegment]) -> str:
    """Export transcript to JSON format."""
    data = {
        "id": transcript.id,
        "recording_id": transcript.recording_id,
        "language": transcript.language,
        "full_text": transcript.full_text,
        "word_count": transcript.word_count,
        "created_at": transcript.created_at.isoformat(),
        "segments": [
            {
                "id": seg.id,
                "index": seg.segment_index,
                "speaker_id": seg.speaker_id,
                "speaker_name": seg.speaker_name,
                "text": seg.text,
                "start_time": seg.start_time,
                "end_time": seg.end_time,
                "confidence": seg.confidence,
            }
            for seg in segments
        ],
    }
    return json.dumps(data, indent=2)


def export_to_docx(
    recording_name: str,
    segments: List[TranscriptSegment],
    include_speakers: bool,
    include_timestamps: bool,
) -> bytes:
    """Export transcript to DOCX format."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title = doc.add_heading(recording_name, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add creation date
    doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_paragraph()

    # Group segments by speaker for cleaner output
    current_speaker = None
    for seg in segments:
        speaker_id = seg.speaker_id or "Unknown"
        speaker_name = seg.speaker_name or speaker_id.replace("speaker_", "Speaker ")

        # Add speaker heading if different from previous
        if include_speakers and speaker_id != current_speaker:
            p = doc.add_paragraph()
            speaker_run = p.add_run(speaker_name)
            speaker_run.bold = True
            speaker_run.font.size = Pt(11)
            speaker_run.font.color.rgb = RGBColor(59, 130, 246)  # Blue
            current_speaker = speaker_id

        # Add segment text
        p = doc.add_paragraph()
        if include_timestamps:
            time_run = p.add_run(f"[{format_vtt_time(seg.start_time)}] ")
            time_run.font.size = Pt(9)
            time_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray

        text_run = p.add_run(seg.text)
        text_run.font.size = Pt(11)

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_to_pdf(
    recording_name: str,
    segments: List[TranscriptSegment],
    include_speakers: bool,
    include_timestamps: bool,
) -> bytes:
    """Export transcript to PDF format."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.75*inch, bottomMargin=0.75*inch)
    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        'TitleStyle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=12,
        alignment=1,  # Center
    )
    speaker_style = ParagraphStyle(
        'SpeakerStyle',
        parent=styles['Normal'],
        fontSize=11,
        textColor=HexColor('#3B82F6'),
        fontName='Helvetica-Bold',
        spaceBefore=12,
        spaceAfter=4,
    )
    text_style = ParagraphStyle(
        'TextStyle',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=8,
        leading=14,
    )
    timestamp_style = ParagraphStyle(
        'TimestampStyle',
        parent=styles['Normal'],
        fontSize=8,
        textColor=HexColor('#808080'),
    )

    story = []

    # Title
    story.append(Paragraph(recording_name, title_style))
    story.append(Paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}", styles['Normal']))
    story.append(Spacer(1, 0.25*inch))

    # Content
    current_speaker = None
    for seg in segments:
        speaker_id = seg.speaker_id or "Unknown"
        speaker_name = seg.speaker_name or speaker_id.replace("speaker_", "Speaker ")

        # Add speaker heading if different from previous
        if include_speakers and speaker_id != current_speaker:
            story.append(Paragraph(speaker_name, speaker_style))
            current_speaker = speaker_id

        # Add segment text
        text = seg.text
        if include_timestamps:
            timestamp = f"<font color='#808080' size='8'>[{format_vtt_time(seg.start_time)}]</font> "
            text = timestamp + text

        story.append(Paragraph(text, text_style))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


@router.post("/{recording_id}/export")
async def export_transcript(
    recording_id: str,
    export_options: ExportRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Generate export in requested format."""
    # Verify access
    recording = await verify_recording_access(db, recording_id, current_user.id)

    if not recording.transcript:
        raise HTTPException(status_code=404, detail="No transcript found for this recording")

    # Get transcript and segments
    transcript = recording.transcript
    result = await db.execute(
        select(TranscriptSegment)
        .where(TranscriptSegment.transcript_id == transcript.id)
        .order_by(TranscriptSegment.segment_index)
    )
    segments = list(result.scalars().all())

    format_type = export_options.format.lower()
    include_speakers = export_options.include_speaker_labels
    include_timestamps = export_options.include_timestamps

    # Generate export based on format
    if format_type == "txt":
        content = export_to_txt(segments, include_speakers, include_timestamps)
        return Response(
            content=content,
            media_type="text/plain",
            headers={"Content-Disposition": f'attachment; filename="{recording.name}.txt"'},
        )

    elif format_type == "srt":
        content = export_to_srt(segments, include_speakers)
        return Response(
            content=content,
            media_type="application/x-subrip",
            headers={"Content-Disposition": f'attachment; filename="{recording.name}.srt"'},
        )

    elif format_type == "vtt":
        content = export_to_vtt(segments, include_speakers)
        return Response(
            content=content,
            media_type="text/vtt",
            headers={"Content-Disposition": f'attachment; filename="{recording.name}.vtt"'},
        )

    elif format_type == "json":
        content = export_to_json(transcript, segments)
        return Response(
            content=content,
            media_type="application/json",
            headers={"Content-Disposition": f'attachment; filename="{recording.name}.json"'},
        )

    elif format_type == "docx":
        content = export_to_docx(recording.name, segments, include_speakers, include_timestamps)
        return Response(
            content=content,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{recording.name}.docx"'},
        )

    elif format_type == "pdf":
        content = export_to_pdf(recording.name, segments, include_speakers, include_timestamps)
        return Response(
            content=content,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{recording.name}.pdf"'},
        )

    else:
        raise HTTPException(status_code=400, detail=f"Unsupported format: {format_options.format}")


@router.get("/{recording_id}/speakers")
async def get_speakers(
    recording_id: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Get speakers for transcript."""
    # Verify access
    recording = await verify_recording_access(db, recording_id, current_user.id)

    if not recording.transcript:
        raise HTTPException(status_code=404, detail="No transcript found")

    # Get unique speakers from segments
    result = await db.execute(
        select(TranscriptSegment.speaker_id, TranscriptSegment.speaker_name)
        .where(TranscriptSegment.transcript_id == recording.transcript.id)
        .distinct()
    )
    speakers_data = result.all()

    speakers = []
    for idx, (speaker_id, speaker_name) in enumerate(speakers_data):
        speakers.append(
            SpeakerResponse(
                id=speaker_id or f"speaker_{idx}",
                name=speaker_name,
                color=get_speaker_color(idx),
            )
        )

    return speakers


@router.post("/{recording_id}/summarize")
async def summarize_transcript(
    recording_id: str,
    summary_type: str = "brief",
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Generate AI summary."""
    # TODO: Implement AI summarization via Ollama
    raise HTTPException(status_code=501, detail="Summarization not yet implemented")


@router.post("/search")
async def search_transcripts(
    query: str,
    project_ids: List[str] | None = None,
    semantic: bool = False,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Search transcripts."""
    # TODO: Implement search
    raise HTTPException(status_code=501, detail="Search not yet implemented")
